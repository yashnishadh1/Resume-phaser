import os
import json
import random
from faker import Faker
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import docx

fake = Faker()

OUTPUT_DIR = "benchmark_dataset"
RESUMES_DIR = os.path.join(OUTPUT_DIR, "resumes")

os.makedirs(RESUMES_DIR, exist_ok=True)

SKILLS_POOL = [
    "Python", "JavaScript", "TypeScript", "Java", "C++", "React", "Angular",
    "SQL", "PostgreSQL", "MongoDB", "AWS", "Docker", "Kubernetes", "Linux",
    "Machine Learning", "Data Analysis", "Communication", "Problem Solving"
]

def generate_pdf(filename, data, template_id):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica", 12)
    y = height - 50

    if template_id == 1:
        # Standard Linear
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, data["name"])
        y -= 20
        c.setFont("Helvetica", 10)
        c.drawString(50, y, f"Email: {data['email']} | Phone: {data['phone']}")
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Skills")
        y -= 20
        c.setFont("Helvetica", 12)
        c.drawString(50, y, ", ".join(data["skills"]))
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Education")
        y -= 20
        c.setFont("Helvetica", 12)
        c.drawString(50, y, data["education"])
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Experience")
        y -= 20
        c.setFont("Helvetica", 12)
        c.drawString(50, y, data["experience"])
    elif template_id == 2:
        # Side by Side Header
        c.setFont("Helvetica-Bold", 18)
        c.drawString(50, y, data["name"])
        c.setFont("Helvetica", 10)
        c.drawRightString(width - 50, y, data["email"])
        y -= 15
        c.drawRightString(width - 50, y, data["phone"])
        y -= 40
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, "Professional Skills")
        y -= 20
        c.setFont("Helvetica", 12)
        for skill in data["skills"]:
            c.drawString(60, y, f"- {skill}")
            y -= 15
        y -= 20
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, "Academic Background")
        y -= 20
        c.setFont("Helvetica", 12)
        c.drawString(50, y, data["education"])
        y -= 40
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, "Work History")
        y -= 20
        c.setFont("Helvetica", 12)
        c.drawString(50, y, data["experience"])
        
    c.save()

def generate_docx(filename, data, template_id):
    doc = docx.Document()
    
    if template_id == 1:
        doc.add_heading(data["name"], 0)
        doc.add_paragraph(f"{data['email']} | {data['phone']}")
        
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(", ".join(data["skills"]))
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph(data["education"])
        
        doc.add_heading('Experience', level=1)
        doc.add_paragraph(data["experience"])
    elif template_id == 2:
        # Table layout
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = data["name"]
        hdr_cells[1].text = f"Email: {data['email']}\nPhone: {data['phone']}"
        
        doc.add_heading('Core Competencies', level=2)
        for skill in data["skills"]:
            doc.add_paragraph(skill, style='List Bullet')
            
        doc.add_heading('Education', level=2)
        doc.add_paragraph(data["education"])
        
        doc.add_heading('Work Experience', level=2)
        doc.add_paragraph(data["experience"])

    doc.save(filename)

def main():
    ground_truth = {}
    
    for i in range(100):
        # Generate Ground Truth Data
        name = fake.name()
        email = fake.email()
        # Ensure standard phone to test parser Regex: '123-456-7890' or '(123) 456-7890'
        phone_format = random.choice(['###-###-####', '(###) ###-####'])
        phone = fake.numerify(phone_format)
        
        num_skills = random.randint(3, 8)
        skills = random.sample(SKILLS_POOL, k=num_skills)
        
        education = f"B.S. in Computer Science, {fake.company()} University"
        experience = f"Software Engineer at {fake.company()} for {random.randint(1, 10)} years"
        
        data = {
            "name": name,
            "email": email,
            "phone": phone,
            "skills": skills,
            "education": education,
            "experience": experience
        }
        
        is_pdf = (i % 2 == 0)
        template_id = random.choice([1, 2])
        
        if is_pdf:
            filename = f"resume_{i}.pdf"
            filepath = os.path.join(RESUMES_DIR, filename)
            generate_pdf(filepath, data, template_id)
        else:
            filename = f"resume_{i}.docx"
            filepath = os.path.join(RESUMES_DIR, filename)
            generate_docx(filepath, data, template_id)
            
        ground_truth[filename] = data
        
    with open(os.path.join(OUTPUT_DIR, "ground_truth.json"), "w") as f:
        json.dump(ground_truth, f, indent=4)
        
    print(f"Generated 100 resumes in {RESUMES_DIR}")
    print("Ground truth saved.")

if __name__ == "__main__":
    main()
