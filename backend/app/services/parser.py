from sqlalchemy.orm import Session
from app.models.resume import Resume, Candidate, CandidateSkill, CandidateExperience, CandidateEducation
import os
import re
from datetime import datetime

try:
    import fitz # PyMuPDF
except ImportError:
    fitz = None
import docx
import pytesseract
try:
    import spacy
except ImportError:
    spacy = None

# pyrefly: ignore [missing-import]
from PIL import Image
import io

# Load NLP model lazily or fallback
nlp = None
if spacy:
    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        import subprocess
        import sys
        try:
            subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
            nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            print(f"Failed to download spaCy model: {e}")

class ResumeParserService:
    def __init__(self, db: Session):
        self.db = db

    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            if fitz:
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text()
                    
                    if not text.strip():
                        for page in doc:
                            pix = page.get_pixmap()
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            text += pytesseract.image_to_string(img)
            else:
                return "Mock PDF text because PyMuPDF is not installed."
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
        return text

    def process_resume(self, resume_id: int):
        resume = self.db.query(Resume).filter(Resume.id == resume_id).first()
        if not resume:
            return

        resume.status = "processing"
        self.db.commit()

        try:
            ext = os.path.splitext(resume.file_path)[1].lower()
            text = ""
            if ext == ".pdf":
                text = self.extract_text_from_pdf(resume.file_path)
            elif ext in [".doc", ".docx"]:
                text = self.extract_text_from_docx(resume.file_path)
            
            # --- NLP EXTRACTION START ---
            
            email_match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text)
            email = email_match.group(0) if email_match else None
            
            phone_match = re.search(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
            phone = phone_match.group(0) if phone_match else None
            
            name = None
            blacklist_headers = {"Core Competencies", "Skills", "Experience", "Summary", "Professional Skills", "Education", "Work History", "Academic Background", "Contact"}
            
            # 1. SpaCy NER for Name
            if nlp:
                doc = nlp(text[:2000]) # Search top part
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        clean_name = ent.text.split('\n')[0].strip()
                        if not any(header.lower() in clean_name.lower() for header in blacklist_headers) and len(clean_name.split()) >= 2:
                            name = clean_name.title()
                            break

            # Regex fallback
            if not name:
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                for line in lines:
                    words = line.split()
                    if 1 <= len(words) <= 4:
                        if not any(char.isdigit() for char in line) and not any(p in line for p in ['.', ',', ':', ';', '(', ')']):
                            if line.lower() not in ['resume', 'cv', 'curriculum vitae'] and not any(header.lower() in line.lower() for header in blacklist_headers):
                                name = line.title()
                                break
            
            if not name:
                name = resume.filename.split('.')[0].replace('_', ' ').replace('-', ' ').title() if resume.filename else "Unknown Candidate"
            
            candidate = Candidate(
                resume_id=resume.id,
                full_name=name,
                email=email,
                phone=phone,
                summary=text[:500] + "..." if len(text) > 500 else text,
                experience_years=0,
                match_score=0
            )
            self.db.add(candidate)
            self.db.commit()
            self.db.refresh(candidate)
            
            # 2. Strict Skills Extraction
            known_skills = [
                # Programming Languages
                "python", "javascript", "typescript", "java", "c", "c++", "c#", "ruby", "go", "php", "swift", "kotlin",
                "rust", "dart", "scala", "perl", "r programming", "r", "shell", "bash", "powershell", "lua", "haskell",
                "clojure", "elixir", "erlang", "objective-c", "assembly",

                # Web Technologies & Frameworks (Frontend)
                "html", "html5", "css", "css3", "react", "react.js", "reactjs", "angular", "vue", "vue.js", "next.js", "nextjs",
                "svelte", "solidjs", "tailwind", "tailwindcss", "bootstrap", "sass", "less", "jquery", "webpack", "vite",
                "redux", "mobx", "zustand", "material-ui", "chakra ui",

                # Backend Frameworks
                "node.js", "nodejs", "node", "django", "flask", "fastapi", "spring", "spring boot", "express", "express.js",
                "nestjs", "laravel", "symfony", "asp.net", ".net", "gin", "echo", "fiber", "ruby on rails", "rails", "sinatra",

                # Mobile Development
                "react native", "flutter", "ios", "android", "xcode", "android studio", "ionic", "xamarin",

                # Databases & Data Management
                "sql", "mysql", "postgresql", "postgres", "mongodb", "redis", "elasticsearch", "sqlite", "oracle", "mariadb",
                "cassandra", "dynamodb", "couchbase", "neo4j", "snowflake", "bigquery", "redshift", "hadoop", "spark", "kafka",
                "graphql", "rest", "rest api", "soap", "grpc", "rabbitmq", "activemq",

                # Cloud & DevOps
                "aws", "amazon web services", "gcp", "google cloud", "azure", "docker", "kubernetes", "linux", "terraform",
                "digitalocean", "heroku", "vercel", "netlify", "jenkins", "github actions", "gitlab ci", "circleci",
                "ansible", "puppet", "chef", "vagrant", "prometheus", "grafana", "datadog", "splunk", "nginx", "apache",
                "git", "bitbucket", "docker compose",

                # AI, ML, Data Science & Analytics
                "machine learning", "nlp", "tensorflow", "pytorch", "pandas", "data analysis", "data science",
                "deep learning", "artificial intelligence", "generative ai", "llms", "openai", "langchain", "hugging face",
                "scikit-learn", "keras", "numpy", "scipy", "matplotlib", "seaborn", "computer vision", "opencv", "tableau", "power bi",

                # Soft Skills & Management
                "communication", "problem solving", "leadership", "agile", "scrum", "teamwork", "critical thinking",
                "time management", "project management", "adaptability", "presentation", "customer service", "collaboration",
                "decision making", "emotional intelligence", "mentoring", "kanban",

                # Basic Tools & Concepts
                "microsoft word", "microsoft excel", "microsoft powerpoint", "ms office", "jira", "trello", "asana", "confluence",
                "postman", "swagger", "figma", "adobe xd", "photoshop", "illustrator",
                
                # Emerging & Other Technologies
                "microservices", "serverless", "blockchain", "web3", "solidity", "smart contracts", "iot", "internet of things",
                "cybersecurity", "ethical hacking", "cryptography", "agile methodology",

                # Business, Finance & Management (MBA, BBA, BCOM)
                "business administration", "marketing", "finance", "accounting", "human resources", "hr", "operations management",
                "supply chain", "logistics", "corporate finance", "taxation", "auditing", "financial modeling", "market research",
                "sales", "retail management", "business strategy", "strategic management", "economics", "cost accounting", "wealth management",

                # Hotel Management (HM)
                "hospitality", "hotel management", "front office", "food & beverage", "f&b", "housekeeping", "event management",
                "catering", "culinary arts", "guest relations", "revenue management", "banquet", "reservation",

                # Law (LLB)
                "legal research", "litigation", "corporate law", "criminal law", "drafting", "legal compliance", "arbitration",
                "intellectual property", "ipr", "contract law", "civil law", "family law", "constitutional law", "paralegal", "advocacy",

                # Arts, Humanities & Sciences (BA, BSC, BCA)
                "sociology", "psychology", "political science", "history", "geography", "literature", "public administration",
                "journalism", "mass communication", "content writing", "editing", "translation", "copywriting",
                "mathematics", "physics", "chemistry", "biology", "statistics", "software engineering", "systems analysis", "network administration",

                # Chemical Fields
                "chemical engineering", "process engineering", "thermodynamics", "mass transfer", "heat transfer", "organic chemistry",
                "inorganic chemistry", "analytical chemistry", "polymer science", "materials science", "spectroscopy", "chromatography",
                "safety engineering", "hazard analysis", "chemical safety", "petrochemistry"
            ]
            
            text_lower = text.lower()
            found_skills = []
            for skill in known_skills:
                # Word boundary strict matching for special characters
                if re.search(rf"(?<![\w]){re.escape(skill)}(?![\w])", text_lower):
                    found_skills.append(skill.title())
                    
            for skill in set(found_skills):
                self.db.add(CandidateSkill(candidate_id=candidate.id, name=skill, type="technical"))

            # 3. Education Extraction
            degrees = ["B.Tech", "B.E.", "BCA", "MCA", "M.Tech", "MBA", "BSc", "MSc", "Diploma", "B.S."]
            extracted_degrees = []
            
            # Simple keyword scan combined with line context
            lines = text.split('\n')
            for i, line in enumerate(lines):
                line_lower = line.lower()
                for deg in degrees:
                    # Use lookaround/boundaries that work with punctuation
                    if re.search(rf"(?:^|\s){re.escape(deg.lower())}(?:\s|$|,)", line_lower):
                        extracted_degrees.append((deg, i))
                        break

            for deg, line_idx in extracted_degrees:
                context_block = " ".join(lines[max(0, line_idx-2):min(len(lines), line_idx+3)])
                
                institution = "Unknown Institution"
                if nlp:
                    doc = nlp(context_block)
                    valid_org_keywords = ["university", "institute", "college", "school", "academy", "iit", "nit", "bits", "mit"]
                    orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG" and any(k in ent.text.lower() for k in valid_org_keywords)]
                    if orgs:
                        institution = orgs[0]

                year_match = re.search(r'\b20[0-2][0-9]\b', context_block)
                year = year_match.group(0) if year_match else "N/A"
                
                cgpa_match = re.search(r'(?:GPA|CGPA|gpa):?\s*([0-9]\.[0-9]{1,2})|([0-9]{1,2}\.[0-9]{1,2})%', context_block)
                cgpa = cgpa_match.group(1) or cgpa_match.group(2) if cgpa_match else None
                
                self.db.add(CandidateEducation(
                    candidate_id=candidate.id,
                    degree=deg,
                    institution=institution,
                    year=year,
                    cgpa=cgpa
                ))

            # 4. Experience Extraction
            experience_section = ""
            exp_match = re.search(r'(?i)(?:Experience|Work History|Employment).*?(?=Education|Skills|Projects|Certifications|$)', text, re.DOTALL)
            if exp_match:
                experience_section = exp_match.group(0)
            else:
                experience_section = text
            
            if nlp:
                doc = nlp(experience_section)
                orgs = []
                dates = []
                
                for ent in doc.ents:
                    if ent.label_ == "ORG" and "university" not in ent.text.lower() and "college" not in ent.text.lower():
                        orgs.append(ent.text)
                    elif ent.label_ == "DATE":
                        dates.append(ent.text)
                
                # De-duplicate ORGs preserving order
                seen = set()
                orgs_unique = [x for x in orgs if not (x in seen or seen.add(x))]
                
                job_titles = ["Engineer", "Developer", "Manager", "Intern", "Analyst", "Consultant", "Architect", "Lead", "Specialist"]
                
                total_years = 0
                for date_str in dates[:3]:
                    years = [int(y) for y in re.findall(r'\b(?:19|20)\d{2}\b', date_str)]
                    if len(years) >= 2:
                        total_years += max(1, years[-1] - years[0])
                    elif len(years) == 1:
                        if "present" in date_str.lower() or "current" in date_str.lower() or "now" in date_str.lower():
                            total_years += max(1, datetime.now().year - years[0])
                        else:
                            total_years += 1

                candidate.experience_years = total_years

                for idx, org in enumerate(orgs_unique[:3]): # Cap at top 3 experiences
                    duration = dates[idx] if idx < len(dates) else "Unknown"
                    position = "Software Engineer"
                    
                    org_idx = experience_section.find(org)
                    if org_idx != -1:
                        start_idx = max(0, org_idx - 100)
                        end_idx = min(len(experience_section), org_idx + len(org) + 100)
                        context_window = experience_section[start_idx:end_idx]
                    else:
                        context_window = experience_section
                    
                    for title in job_titles:
                        if re.search(rf"(?i)\b{title}\b", context_window):
                            position = title
                            break
                            
                    self.db.add(CandidateExperience(
                        candidate_id=candidate.id,
                        company=org,
                        position=position,
                        duration=duration,
                        description=""
                    ))

            resume.status = "completed"
            self.db.commit()

        except Exception as e:
            resume.status = "error"
            self.db.commit()
            print(f"Error parsing resume {resume_id}: {e}")
