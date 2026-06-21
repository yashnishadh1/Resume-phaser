from docx import Document

def create_resume():
    doc = Document()
    doc.add_heading('Taran', 0)
    doc.add_paragraph('Software Engineer with 5 years of experience in Python and React.')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, JavaScript, React, FastAPI, SQL')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Developer at Tech Corp (2020-Present)')
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. Computer Science')
    doc.save('taran_resume.docx')

if __name__ == '__main__':
    create_resume()
    print("Created taran_resume.docx successfully.")
